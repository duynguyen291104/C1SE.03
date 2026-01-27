"""
Module xuất đề kiểm tra ra file DOCX
"""
from pathlib import Path
from loguru import logger
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import Exam, ExamMatrix, Blueprint, ExportRequest


class DOCXExporter:
    """Xuất đề kiểm tra ra file DOCX"""
    
    def __init__(self, template_path: str = None):
        """
        Args:
            template_path: Đường dẫn template DOCX (nếu có)
        """
        self.template_path = template_path
    
    def export(self, request: ExportRequest, output_path: str) -> str:
        """
        Xuất đề kiểm tra ra DOCX
        
        Args:
            request: Export request
            output_path: Đường dẫn file output
            
        Returns:
            Đường dẫn file đã xuất
        """
        logger.info(f"📄 Exporting exam to DOCX: {output_path}")
        
        # Create document
        if self.template_path and Path(self.template_path).exists():
            doc = DocxDocument(self.template_path)
        else:
            doc = DocxDocument()
        
        # Add content
        self._add_header(doc, request.exam)
        self._add_matrix(doc, request.matrix, request.blueprint)
        self._add_specification(doc, request.matrix, request.blueprint)
        doc.add_page_break()
        self._add_exam_questions(doc, request.exam)
        
        if request.include_answer_key:
            doc.add_page_break()
            self._add_answer_key(doc, request.exam, request.include_rubric)
        
        # Save
        doc.save(output_path)
        logger.info(f"✅ Exported to: {output_path}")
        
        return output_path
    
    def _add_header(self, doc: DocxDocument, exam: Exam):
        """Thêm trang bìa"""
        # School info (placeholder)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TRƯỜNG THCS/THPT ...\n")
        run.font.size = Pt(12)
        run.font.bold = True
        
        run = p.add_run("TỔ: ...\n")
        run.font.size = Pt(12)
        
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(exam.title.upper())
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Subject & grade
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subject_text = f"Môn: {exam.subject}"
        if exam.grade:
            subject_text += f" - Lớp {exam.grade}"
        run = p.add_run(subject_text)
        run.font.size = Pt(13)
        run.font.italic = True
        
        # Time
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Thời gian: {exam.time_minutes} phút (không kể thời gian phát đề)")
        run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
    
    def _add_matrix(self, doc: DocxDocument, matrix: ExamMatrix, blueprint: Blueprint):
        """Thêm ma trận đề (Phụ lục 1)"""
        # Title
        p = doc.add_paragraph()
        run = p.add_run("PHỤ LỤC 1: MA TRẬN ĐỀ KIỂM TRA")
        run.font.size = Pt(13)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table
        n_cols = 8  # Chủ đề, Biết, Hiểu, VD, VDC, Tổng câu, Tổng điểm, Tỷ lệ
        n_rows = len(matrix.items_plan) + 2  # +1 header, +1 total
        
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = 'Table Grid'
        
        # Header row
        headers = ["Chủ đề", "Biết", "Hiểu", "Vận dụng", "VD cao", "Tổng câu", "Tổng điểm", "Tỷ lệ %"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_bold(cell)
        
        # Data rows
        topic_stats = {}
        for row_idx, item in enumerate(matrix.items_plan, start=1):
            topic = next((t for t in blueprint.topics if t.topic_id == item.topic_id), None)
            topic_name = topic.name if topic else item.topic_id
            
            # Aggregate by topic
            if topic_name not in topic_stats:
                topic_stats[topic_name] = {"biet": 0, "hieu": 0, "vandung": 0, "vandungcao": 0, "total": 0, "points": 0}
            
            topic_stats[topic_name][item.cognitive_level] += item.n_questions
            topic_stats[topic_name]["total"] += item.n_questions
            topic_stats[topic_name]["points"] += item.n_questions * item.points_each
        
        # Fill table
        row_idx = 1
        for topic_name, stats in topic_stats.items():
            table.rows[row_idx].cells[0].text = topic_name
            table.rows[row_idx].cells[1].text = str(stats["biet"])
            table.rows[row_idx].cells[2].text = str(stats["hieu"])
            table.rows[row_idx].cells[3].text = str(stats["vandung"])
            table.rows[row_idx].cells[4].text = str(stats["vandungcao"])
            table.rows[row_idx].cells[5].text = str(stats["total"])
            table.rows[row_idx].cells[6].text = f"{stats['points']:.1f}"
            table.rows[row_idx].cells[7].text = f"{stats['points'] / matrix.global_config.total_points * 100:.0f}"
            row_idx += 1
        
        # Total row
        total_biet = sum(s["biet"] for s in topic_stats.values())
        total_hieu = sum(s["hieu"] for s in topic_stats.values())
        total_vandung = sum(s["vandung"] for s in topic_stats.values())
        total_vandungcao = sum(s["vandungcao"] for s in topic_stats.values())
        total_questions = sum(s["total"] for s in topic_stats.values())
        total_points = sum(s["points"] for s in topic_stats.values())
        
        table.rows[-1].cells[0].text = "TỔNG"
        table.rows[-1].cells[1].text = str(total_biet)
        table.rows[-1].cells[2].text = str(total_hieu)
        table.rows[-1].cells[3].text = str(total_vandung)
        table.rows[-1].cells[4].text = str(total_vandungcao)
        table.rows[-1].cells[5].text = str(total_questions)
        table.rows[-1].cells[6].text = f"{total_points:.1f}"
        table.rows[-1].cells[7].text = "100"
        self._set_cell_bold(table.rows[-1].cells[0])
        
        doc.add_paragraph()  # Spacing
    
    def _add_specification(self, doc: DocxDocument, matrix: ExamMatrix, blueprint: Blueprint):
        """Thêm bảng đặc tả"""
        # Title
        p = doc.add_paragraph()
        run = p.add_run("BẢNG ĐẶC TẢ")
        run.font.size = Pt(13)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table
        table = doc.add_table(rows=len(matrix.items_plan) + 1, cols=6)
        table.style = 'Table Grid'
        
        # Header
        headers = ["STT", "Chủ đề", "Mức độ", "Dạng câu", "Số câu", "Điểm"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_bold(cell)
        
        # Data
        cognitive_map = {
            "biet": "Biết",
            "hieu": "Hiểu",
            "vandung": "Vận dụng",
            "vandungcao": "Vận dụng cao"
        }
        
        type_map = {
            "mcq_single": "Trắc nghiệm",
            "mcq_multiple": "Trắc nghiệm nhiều đáp án",
            "true_false": "Đúng/Sai",
            "fill_blank": "Điền khuyết",
            "short_answer": "Tự luận ngắn",
            "essay": "Tự luận"
        }
        
        for idx, item in enumerate(matrix.items_plan, start=1):
            topic = next((t for t in blueprint.topics if t.topic_id == item.topic_id), None)
            topic_name = topic.name if topic else item.topic_id
            
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = topic_name
            table.rows[idx].cells[2].text = cognitive_map.get(item.cognitive_level, item.cognitive_level)
            table.rows[idx].cells[3].text = type_map.get(item.type, item.type)
            table.rows[idx].cells[4].text = str(item.n_questions)
            table.rows[idx].cells[5].text = f"{item.points_each:.2f}"
        
        doc.add_paragraph()
    
    def _add_exam_questions(self, doc: DocxDocument, exam: Exam):
        """Thêm đề kiểm tra"""
        # Title
        p = doc.add_paragraph()
        run = p.add_run("ĐỀ BÀI")
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Questions
        for q in exam.questions:
            # Question stem
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {q.id.replace('Q', '')}: ")
            run.font.bold = True
            run = p.add_run(q.stem)
            
            # Options (for MCQ)
            if q.options:
                for option in q.options:
                    p = doc.add_paragraph(option, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
            
            doc.add_paragraph()  # Spacing
    
    def _add_answer_key(self, doc: DocxDocument, exam: Exam, include_rubric: bool):
        """Thêm đáp án và hướng dẫn chấm"""
        # Title
        p = doc.add_paragraph()
        run = p.add_run("ĐÁP ÁN VÀ HƯỚNG DẪN CHẤM")
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Answer table
        table = doc.add_table(rows=len(exam.questions) + 1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        headers = ["Câu", "Đáp án", "Điểm", "Giải thích"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_bold(cell)
        
        # Data
        for idx, q in enumerate(exam.questions, start=1):
            table.rows[idx].cells[0].text = q.id.replace('Q', '')
            table.rows[idx].cells[1].text = q.answer
            table.rows[idx].cells[2].text = f"{q.points:.2f}"
            table.rows[idx].cells[3].text = q.explanation or ""
        
        # Rubric for essay questions
        if include_rubric:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("CHI TIẾT CHẤM CÂU TỰ LUẬN")
            run.font.bold = True
            
            for q in exam.questions:
                if q.rubric:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    run = p.add_run(f"Câu {q.id.replace('Q', '')}: ")
                    run.font.bold = True
                    
                    for criterion in q.rubric.criteria:
                        p = doc.add_paragraph(
                            f"• {criterion.get('description', '')}: {criterion.get('points', 0)} điểm",
                            style='List Bullet'
                        )
    
    def _set_cell_bold(self, cell):
        """Set cell text bold"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
